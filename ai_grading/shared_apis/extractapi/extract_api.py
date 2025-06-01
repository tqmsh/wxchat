import zipfile
import tarfile
import shutil
#import fitz  # PyMuPDF
from PyPDF2 import PdfReader, PdfWriter
import re
import os
from abc import ABC, abstractmethod
from typing import List, Dict, Any
from pdf2image import convert_from_path
from PIL import Image
from docx import Document
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE
import requests
from io import BytesIO

class FileProcessor(ABC):
    """
    Abstract base class for processing files.

    Attributes:
        input_path (str): Path to the input file or directory.
        output_path (str): Path to the output directory.
    """
    def __init__(self, input_path, output_path):
        """
        Initializes the FileProcessor with input and output paths.

        Args:
            input_path (str): Path to the input file or directory.
            output_path (str): Path to the output directory.
        """
        self.input_path = input_path
        self.output_path = output_path

    @abstractmethod
    def process(self):
        """
        Abstract method for processing files. Must be implemented by subclasses.
        """
        pass

    def ensure_output_directory(self):
        """
        Ensures that the output directory exists. Creates the directory if it does not exist.

        Returns:
            str: Error message if directory creation fails, otherwise None.
        """
        if not os.path.isdir(self.output_path):
            try:
                os.makedirs(self.output_path)
            except OSError as e:
                return f"Error: Could not create output directory {self.output_path}. {e}"
        return None

class ZipExtractor(FileProcessor):
    """
    Extracts the contents of a ZIP file to the specified output directory.

    Attributes:
        input_path (str): Path to the ZIP file.
        output_path (str): Path to the output directory where contents will be extracted.
    """
    def process(self):
        """
        Extracts the ZIP file to the output directory.

        Returns:
            str: Error message if extraction fails, otherwise None.
        """
        error = self.ensure_output_directory()
        if error:
            return error

        try:
            with zipfile.ZipFile(self.input_path, 'r') as zip_ref:
                zip_ref.extractall(self.output_path)
            print(f"Contents of '{self.input_path}' have been extracted to '{self.output_path}'.")
            return None
        except FileNotFoundError:
            return f"Error: The file '{self.input_path}' was not found."
        except zipfile.BadZipFile:
            return f"Error: The file '{self.input_path}' is not a zip file or it is corrupted."
        except PermissionError:
            return f"Error: Permission denied for '{self.output_path}'."
        except Exception as e:
            return f"An unexpected error occurred: {str(e)}"

class TarGzExtractor(FileProcessor):
    """
    Extracts the contents of a tar.gz file to the specified output directory.

    Attributes:
        input_path (str): Path to the tar.gz file.
        output_path (str): Path to the output directory where contents will be extracted.
    """
    def process(self):
        """
        Extracts the tar.gz file to the output directory.

        Returns:
            str: Error message if extraction fails, otherwise None.
        """
        if not os.path.isfile(self.input_path):
            return f"Error: The file {self.input_path} does not exist."

        error = self.ensure_output_directory()
        if error:
            return error

        try:
            with tarfile.open(self.input_path, "r:gz") as tar:
                tar.extractall(path=self.output_path)
            return None
        except tarfile.TarError as e:
            return f"Error: Problem with the tar.gz file {self.input_path}. {e}"
        except Exception as e:
            return f"Error: An unexpected error occurred. {e}"

class FileCopier(FileProcessor):
    """
    Copies files with a specific extension from an input directory to an output directory.

    Attributes:
        input_path (str): Path to the input directory.
        output_path (str): Path to the output directory where files will be copied.
        file_extension (str): The file extension of files to be copied.
    """
    def __init__(self, input_path, output_path, file_extension):
        """
        Initializes the FileCopier with input and output paths and file extension.

        Args:
            input_path (str): Path to the input directory.
            output_path (str): Path to the output directory.
            file_extension (str): The file extension of files to be copied.
        """
        super().__init__(input_path, output_path)
        self.file_extension = file_extension

    def process(self):
        """
        Copies files with the specified extension from the input directory to the output directory.

        Returns:
            str: Error message if copying fails, otherwise None.
        """
        if not os.path.isdir(self.input_path):
            return f"Error: Input folder '{self.input_path}' does not exist."

        error = self.ensure_output_directory()
        if error:
            return error

        try:
            for filename in os.listdir(self.input_path):
                if filename.endswith(self.file_extension):
                    src_path = os.path.join(self.input_path, filename)
                    dst_path = os.path.join(self.output_path, filename)
                    shutil.copy2(src_path, dst_path)
            return None
        except OSError as e:
            return f"Error: An error occurred while copying files: {e}"


class PDFToImageConverter(FileProcessor):
    """
    Converts PDF files from an input directory to image files in an output directory.
    Can either combine all pages into a single image or save them as separate images.

    Attributes:
        input_path (str): Path to the input directory.
        output_path (str): Path to the output directory where images will be saved.
        combine_images (bool): Whether to combine all pages into a single image or save them separately.
    """

    def __init__(self, input_path, output_path, combine_images=True):
        """
        Initializes the PDFToImageConverter with input and output paths, and combining option.

        Args:
            input_path (str): Path to the input directory.
            output_path (str): Path to the output directory.
            combine_images (bool): Whether to combine all pages into a single image. Defaults to True.
        """
        super().__init__(input_path, output_path)
        self.combine_images = combine_images

    def process(self):
        """
        Converts PDF files to images and saves them to the output directory.
        If combine_images is True, saves a single combined image.
        If combine_images is False, saves separate images for each page.

        Returns:
            str: Error message if processing fails, otherwise None.
        """
        #If it isn't a directory but it does exist, it must be a file, so let's make sure it's a pdf and then go for it
        try:
          if not os.path.isdir(self.input_path) and os.path.exists(self.input_path):
              print(self.output_path)
              if self.input_path.endswith(".pdf"):
                  pdf_path = self.input_path
                  pages = convert_from_path(pdf_path, 500)
  
                  if self.combine_images:
                      self._save_combined_image(pages, os.path.basename(self.input_path))
                  else:
                      self._save_separate_images(pages, os.path.basename(self.input_path))
                  return None #It worked!
              return f"Error: input file is not a pdf nor a directory."
        except Exception as e:
          return f"Error: An error occurred while converting PDF files: {str(e)}"
        error = self.ensure_output_directory()
        if error:
            return error

        try:
            for file in os.listdir(self.input_path):
                if file.endswith(".pdf"):
                    pdf_path = os.path.join(self.input_path, file)
                    pages = convert_from_path(pdf_path, 500)

                    if self.combine_images:
                        self._save_combined_image(pages, file)
                    else:
                        self._save_separate_images(pages, file)

            return None
        except Exception as e:
            return f"Error: An error occurred while converting PDF files: {str(e)}"

    def _save_combined_image(self, pages, file_name):
        """
        Saves all pages of a PDF as a single combined image.

        Args:
            pages (list): List of PIL Image objects representing PDF pages.
            file_name (str): Name of the original PDF file.
        """
        width, height = pages[0].size
        total_height = height * len(pages)
        combined_image = Image.new('RGB', (width, total_height))

        y_offset = 0
        for page in pages:
            combined_image.paste(page, (0, y_offset))
            y_offset += height

        output_file = os.path.join(self.output_path, f'{os.path.splitext(file_name)[0]}_combined.jpg')
        combined_image.save(output_file, 'JPEG')

    def _save_separate_images(self, pages, file_name):
        """
        Saves each page of a PDF as a separate image.

        Args:
            pages (list): List of PIL Image objects representing PDF pages.
            file_name (str): Name of the original PDF file.
        """
        base_name = os.path.splitext(file_name)[0]
        for i, page in enumerate(pages, start=1):
            output_file = os.path.join(self.output_path, f'{base_name}_{i}.jpg')
            page.save(output_file, 'JPEG')

class PdfSplitter(FileProcessor):
    """
    Splits a PDF file into multiple files based on a regex pattern that matches text.
    Attributes:
        input_path (str): Path to the PDF file.
        output_path (str): Path to the output directory where split PDFs will be saved.
        regex_pattern (str): Regex pattern to determine split points in the PDF.
        keep_file_name (bool): Flag to indicate if the original file name should be appended to the split files.
    """
    def __init__(self, input_path, output_path, regex_pattern, keep_file_name=True):
        """
        Initializes the PdfSplitter with input and output paths and a regex pattern.
        Args:
            input_path (str): Path to the PDF file.
            output_path (str): Path to the output directory.
            regex_pattern (str): Regex pattern to determine split points.
            keep_file_name (bool): Flag to indicate if the original file name should be appended to the split files.
        """
        super().__init__(input_path, output_path)
        self.regex_pattern = regex_pattern
        self.keep_file_name = keep_file_name

    def process(self):
        """
        Splits the PDF file into multiple files based on the regex pattern. Each split file contains pages
        from the last match to the next match.
        Returns:
            str: Error message if splitting fails, otherwise None.
        """
        if not os.path.isfile(self.input_path):
            return "Input PDF file does not exist."
        
        error = self.ensure_output_directory()
        if error:
            return error

        try:
            pdf_reader = PdfReader(self.input_path)
            pattern = re.compile(self.regex_pattern)
            split_points = [0]  # Start with the first page

            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                if pattern.search(page.extract_text()):
                    split_points.append(page_num)

            split_points.append(len(pdf_reader.pages))  # End with the last page

            if len(split_points) == 2:  # Only start and end points
                return "No matches found for the provided regex pattern."

            # Get the original file name without the extension if keep_file_name is True
            file_name = ""
            if self.keep_file_name:
                file_name = os.path.splitext(os.path.basename(self.input_path))[0]
                file_name += "_"

            for i in range(len(split_points) - 1):
                start_page = split_points[i]
                end_page = split_points[i + 1]
                pdf_writer = PdfWriter()

                for page_num in range(start_page, end_page):
                    pdf_writer.add_page(pdf_reader.pages[page_num])

                output_pdf_path = os.path.join(self.output_path, f"{file_name}split_{i+1}.pdf")
                with open(output_pdf_path, 'wb') as output_file:
                    pdf_writer.write(output_file)

            return None
        except Exception as e:
            return str(e)
       

class PagePdfSplitter(FileProcessor):
    """
    Splits a PDF file into multiple files based on a given number of pages.
    Attributes:
        input_path (str): Path to the PDF file.
        output_path (str): Path to the output directory where split PDFs will be saved.
        pages_per_split (int): Number of pages per split PDF.
        keep_file_name (bool): Flag to indicate if the original file name should be appended to the split files.
    """
    def __init__(self, input_path, output_path, pages_per_split, keep_file_name=True):
        """
        Initializes the PagePdfSplitter with input and output paths and the number of pages per split.
        Args:
            input_path (str): Path to the PDF file.
            output_path (str): Path to the output directory.
            pages_per_split (int): Number of pages per split PDF.
            keep_file_name (bool): Flag to indicate if the original file name should be appended to the split files.
        """
        super().__init__(input_path, output_path)
        self.pages_per_split = pages_per_split
        self.keep_file_name = keep_file_name

    def process(self):
        """
        Splits the PDF file into multiple files based on the number of pages per split.
        Returns:
            str: Error message if splitting fails, otherwise None.
        """
        if not os.path.isfile(self.input_path):
            return "Input PDF file does not exist."
        
        error = self.ensure_output_directory()
        if error:
            return error

        try:
            pdf_reader = PdfReader(self.input_path)
            total_pages = len(pdf_reader.pages)

            # Get the original file name without the extension if keep_file_name is True
            file_name = ""
            if self.keep_file_name:
                file_name = os.path.splitext(os.path.basename(self.input_path))[0]
                file_name += "_"

            split_num = 1
            for start_page in range(0, total_pages, self.pages_per_split):
                end_page = min(start_page + self.pages_per_split, total_pages)
                pdf_writer = PdfWriter()

                for page_num in range(start_page, end_page):
                    pdf_writer.add_page(pdf_reader.pages[page_num])

                output_pdf_path = os.path.join(self.output_path, f"{file_name}split_{split_num}.pdf")
                with open(output_pdf_path, 'wb') as output_file:
                    pdf_writer.write(output_file)

                split_num += 1

            return None
        except Exception as e:
            return str(e)            


class RangePdfSplitter(FileProcessor):
    """
    Extracts a sub-PDF from an inclusive range of pages.
    Attributes:
        input_path (str): Path to the PDF file.
        output_path (str): Path to the output directory where the extracted PDF will be saved.
        start_page (int): The starting page number (1-indexed) for the extraction.
        end_page (int or None): The ending page number (1-indexed) for the extraction. If None, extract to the end of the PDF.
        keep_file_name (bool): Flag to indicate if the original file name should be appended to the output file.
    """
    def __init__(self, input_path, output_path, start_page, end_page=None, keep_file_name=True):
        """
        Initializes the RangePdfSplitter with input and output paths, and the range of pages to extract.
        Args:
            input_path (str): Path to the PDF file.
            output_path (str): Path to the output directory.
            start_page (int): The starting page number (1-indexed) for the extraction.
            end_page (int or None): The ending page number (1-indexed) for the extraction. If None, extract to the end of the PDF.
            keep_file_name (bool): Flag to indicate if the original file name should be appended to the output file.
        """
        super().__init__(input_path, output_path)
        self.start_page = start_page
        self.end_page = end_page
        self.keep_file_name = keep_file_name

    def process(self):
        """
        Extracts the PDF file based on the specified range of pages.
        Returns:
            str: Error message if extraction fails, otherwise None.
        """
        if not os.path.isfile(self.input_path):
            return "Input PDF file does not exist."
        
        error = self.ensure_output_directory()
        if error:
            return error

        try:
            pdf_reader = PdfReader(self.input_path)
            total_pages = len(pdf_reader.pages)

            # Validate start_page and end_page
            if self.start_page < 1 or self.start_page > total_pages:
                return "Invalid start page."
            if self.end_page is not None and (self.end_page < self.start_page or self.end_page > total_pages):
                return "Invalid end page."
            
            # If end_page is None, set it to the last page
            if self.end_page is None:
                self.end_page = total_pages

            # Get the original file name without the extension if keep_file_name is True
            file_name = ""
            if self.keep_file_name:
                file_name = os.path.splitext(os.path.basename(self.input_path))[0]
                file_name += "_"

            pdf_writer = PdfWriter()
            for page_num in range(self.start_page - 1, self.end_page):
                pdf_writer.add_page(pdf_reader.pages[page_num])

            output_pdf_path = os.path.join(self.output_path, f"{file_name}pages_{self.start_page}_to_{self.end_page}.pdf")
            with open(output_pdf_path, 'wb') as output_file:
                pdf_writer.write(output_file)

            return None
        except Exception as e:
            return str(e)

class ImageSplitter(FileProcessor):
    """
    Processes an image by sending it to a FastAPI endpoint, receives subimages, and saves them.
    
    Attributes:
        input_path (str): Path to the input image file.
        output_path (str): Path to the output directory where subimages will be saved.
        api_url (str): URL of the FastAPI endpoint for image processing.
        keep_file_name (bool): Flag to indicate if the original file name should be appended to the subimages.
    """
    
    def __init__(self, input_path, output_path, 
                 api_url="http://ece-nebula11.eng.uwaterloo.ca:8004/process-image/", 
                 keep_file_name=True):
        """
        Initializes the ImageSubdivider with input and output paths and the API URL.
        
        Args:
            input_path (str): Path to the input image file.
            output_path (str): Path to the output directory.
            api_url (str): URL of the FastAPI endpoint for image processing. 
                           Defaults to "http://ece-nebula11.eng.uwaterloo.ca:8004/process-image/".
            keep_file_name (bool): Flag to indicate if the original file name should be appended to the subimages.
        """
        super().__init__(input_path, output_path)
        self.api_url = api_url
        self.keep_file_name = keep_file_name
    
    def process(self):
        """
        Processes the image by sending it to the API, receives subimages, and saves them.
        
        Returns:
            str: Error message if processing fails, otherwise None.
        """
        if not os.path.isfile(self.input_path):
            return "Input image file does not exist."
        
        error = self.ensure_output_directory()
        if error:
            return error
        
        try:
            # Send the image to the FastAPI endpoint
            with open(self.input_path, 'rb') as image_file:
                files = {'file': (os.path.basename(self.input_path), image_file, 'image/jpeg')}
                headers = {'accept': 'image/*'}
                response = requests.post(self.api_url, files=files, headers=headers)
            
            if response.status_code != 200:
                return f"API request failed with status code {response.status_code}"
            
            # Process the received zip file
            zip_content = BytesIO(response.content)
            with zipfile.ZipFile(zip_content) as zip_ref:
                # Get the original file name without the extension if keep_file_name is True
                file_name = ""
                if self.keep_file_name:
                    file_name = os.path.splitext(os.path.basename(self.input_path))[0]
                    file_name += "_"
                
                # Extract and save subimages
                for i, file_info in enumerate(zip_ref.infolist(), start=1):
                    subimage_data = zip_ref.read(file_info.filename)
                    subimage_name = f"{file_name}subimage_{i}{os.path.splitext(file_info.filename)[1]}"
                    subimage_path = os.path.join(self.output_path, subimage_name)
                    
                    with open(subimage_path, 'wb') as subimage_file:
                        subimage_file.write(subimage_data)
            
            return None
        except Exception as e:
            return str(e)

class TextSplitter(FileProcessor):
    """
    Splits a text file into multiple files based on a regex pattern that matches text.
    Attributes:
        input_path (str): Path to the text file.
        output_path (str): Path to the output directory where split text files will be saved.
        regex_pattern (str): Regex pattern to determine split points in the text file.
        keep_file_name (bool): Flag to indicate if the original file name should be appended to the split files.
    """
    def __init__(self, input_path, output_path, regex_pattern, keep_file_name=True):
        """
        Initializes the TextSplitter with input and output paths and a regex pattern.
        Args:
            input_path (str): Path to the text file.
            output_path (str): Path to the output directory.
            regex_pattern (str): Regex pattern to determine split points.
            keep_file_name (bool): Flag to indicate if the original file name should be appended to the split files.
        """
        super().__init__(input_path, output_path)
        self.regex_pattern = regex_pattern
        self.keep_file_name = keep_file_name

    def process(self):
        """
        Splits the text file into multiple files based on the regex pattern. Each split file contains lines
        from the last match to the next match.
        Returns:
            str: Error message if splitting fails, otherwise None.
        """
        if not os.path.isfile(self.input_path):
            return "Input text file does not exist."
        
        error = self.ensure_output_directory()
        if error:
            return error

        try:
            with open(self.input_path, 'r') as file:
                lines = file.readlines()

            pattern = re.compile(self.regex_pattern)
            split_points = []

            for i, line in enumerate(lines):
                if pattern.search(line):
                    split_points.append(i)

            split_points.append(len(lines))  # End with the last line

            if len(split_points) == 1:  # No matches found
                return "No matches found for the provided regex pattern."

            # Get the original file name without the extension if keep_file_name is True
            file_name = ""
            if self.keep_file_name:
                file_name = os.path.splitext(os.path.basename(self.input_path))[0]
                file_name += "_"

            file_counter = 1
            for i in range(len(split_points) - 1):
                start_line = split_points[i]
                end_line = split_points[i + 1]
                
                # Check if this segment is empty (only for the first segment)
                if i == 0 and start_line == end_line:
                    continue

                output_text_path = os.path.join(self.output_path, f"{file_name}split_{file_counter}.txt")
                with open(output_text_path, 'w') as output_file:
                    output_file.writelines(lines[start_line:end_line])
                file_counter += 1

            return None
        except Exception as e:
            return str(e)


class DocxSplitter:
    """
    Splits a docx file into multiple files based on a specified style and regex pattern.
    
    Attributes:
        input_path (str): Path to the input docx file.
        output_path (str): Path to the output directory where split docx files will be saved.
        style_name (str): Name of the style to match (e.g., 'Heading 2').
        regex_pattern (str): Regex pattern that the text with the specified style must match.
        keep_file_name (bool): Flag to indicate if the original file name should be prepended to the split files.
        
        TODO: The explanation of what this does is WAY more ambitious than what it actualy does. It really only splits on
        Heading 1 or Heading 2. It needs to be updated, but it suits my pruposes.
    """

    def __init__(self, input_path, output_path, regex_pattern, style_name="Heading 2", keep_file_name=True):
        """
        Initializes the DocxSplitter with input and output paths, style name, regex pattern, and keep_file_name option.
        
        Args:
            input_path (str): Path to the input docx file.
            output_path (str): Path to the output directory.
            style_name (str): Name of the style to match.
            regex_pattern (str): Regex pattern that the text with the specified style must match.
            keep_file_name (bool): Flag to indicate if the original file name should be prepended to the split files.
        """
        self.input_path = input_path
        self.output_path = output_path
        self.style_name = style_name
        self.regex_pattern = regex_pattern
        self.keep_file_name = keep_file_name

    def ensure_output_directory(self):
        """
        Ensures that the output directory exists.
        
        Returns:
            str: Error message if creating the output directory fails, otherwise None.
        """
        try:
            os.makedirs(self.output_path, exist_ok=True)
            return None
        except Exception as e:
            return str(e)

    def process(self):
        """
        Splits the docx file into multiple files based on the specified style and regex pattern.
        Each split file contains content under each matched heading, including the heading itself.
        
        Returns:
            str: Error message if splitting fails, otherwise None.
        """
        if not os.path.isfile(self.input_path):
            return "Input docx file does not exist."
        
        error = self.ensure_output_directory()
        if error:
            return error
    
        try:
            doc = Document(self.input_path)
            pattern = re.compile(self.regex_pattern)
            split_points = []
            
            # Find all paragraphs that match both the style and regex pattern
            for i, para in enumerate(doc.paragraphs):
                if para.style.name == self.style_name and pattern.search(para.text):
                    split_points.append(i)
            
            if not split_points:
                return "No matches found for the provided style and regex pattern."
            
            # Get the original file name without the extension if keep_file_name is True
            file_name = ""
            if self.keep_file_name:
                file_name = os.path.splitext(os.path.basename(self.input_path))[0]
                file_name += "_"
            
            # Split the document based on the headings
            for i in range(len(split_points)):
                start = split_points[i]  # Start at the match
                end = split_points[i + 1] if i + 1 < len(split_points) else len(doc.paragraphs)
                
                new_doc = Document()
                within_section = False
                
                for para in doc.paragraphs[start:end]:
                    if para.style.name == self.style_name:
                        within_section = True
                    elif within_section and para.style.name in ["Heading 1", "Heading 2"]:
                        break  # Stop when we reach a higher or equal-level heading
                    
                    if within_section:
                        new_para = new_doc.add_paragraph()
                        new_para.text = para.text
                        
                        # Check if the style already exists in the new document
                        if para.style.name in new_doc.styles:
                            new_para.style = new_doc.styles[para.style.name]
                        else:
                            # If the style doesn't exist, add it
                            new_style = new_doc.styles.add_style(para.style.name, WD_STYLE_TYPE.PARAGRAPH)
                            new_style.font.size = para.style.font.size
                            new_para.style = new_style
                
                output_file_path = os.path.join(self.output_path, f"{file_name}split_{i + 1}.docx")
                new_doc.save(output_file_path)
            
            return None
        except Exception as e:
            return str(e)
            
class Converger(FileProcessor):
    """
    Copies all files from subdirectories of the source directory into a single flat target directory.
    Attributes:
        input_path (str): Path to the directory containing subdirectories with files to be copied.
        output_path (str): Path to the target directory where all files will be copied.
    """
    def __init__(self, input_path, output_path):
        """
        Initializes the Converger with input and output paths.
        Args:
            input_path (str): Path to the directory with subdirectories.
            output_path (str): Path to the target directory.
        """
        super().__init__(input_path, output_path)

    def process(self):
        """
        Copies all files from the subdirectories of the input directory into the output directory.
        Returns:
            str: Error message if copying fails, otherwise None.
        """
        if not os.path.isdir(self.input_path):
            return "Input directory does not exist."

        error = self.ensure_output_directory()
        if error:
            return error

        try:
            for subdir in os.listdir(self.input_path):
                subdir_path = os.path.join(self.input_path, subdir)
                if os.path.isdir(subdir_path):
                    for item in os.listdir(subdir_path):
                        item_path = os.path.join(subdir_path, item)
                        if os.path.isfile(item_path):
                            shutil.copy(item_path, self.output_path)
                        elif os.path.isdir(item_path):
                            print(f"Warning: Ignoring nested directory {item_path}")
            return None
        except Exception as e:
            return str(e)

class ProcessingStep:
    def __init__(self, processor_class, mode='flat', **kwargs):
        self.processor_class = processor_class
        self.mode = mode
        self.kwargs = kwargs
        
class Project:
    def __init__(self, name: str, base_input_path: str, base_output_path: str):
        self.name = name
        self.base_input_path = base_input_path
        self.base_output_path = base_output_path
        self.steps: List[ProcessingStep] = []

    def add_step(self, processor_class, mode='flat', **kwargs):
        step_index = len(self.steps) + 1
        expected_output_path = os.path.join(self.base_output_path, f"step_{step_index}")
        self.steps.append(ProcessingStep(processor_class, mode, **kwargs))
        
        return f"{expected_output_path}"
        
    def run(self):
        current_input_path = self.base_input_path
        for i, step in enumerate(self.steps):
            step_output_path = os.path.join(self.base_output_path, f"step_{i+1}")
            os.makedirs(step_output_path, exist_ok=True)

            if step.mode == 'converge':
                # Special handling for converge mode
                processor = step.processor_class(current_input_path, step_output_path, **step.kwargs)
                error = processor.process()
                if error:
                    print(f"Error in step {step.processor_class.__name__}: {error}")
            else:

                if os.path.isfile(current_input_path):
                    self._process_file(step, current_input_path, step_output_path)
                elif os.path.isdir(current_input_path):
                    self._process_directory(step, current_input_path, step_output_path)
                else:
                    print(f"Error: {current_input_path} is neither a file nor a directory.")
                    return
    
            current_input_path = step_output_path

        print(f"Project '{self.name}' completed. Final output is in: {current_input_path}")

    def _process_file(self, step: ProcessingStep, input_path: str, output_path: str):
        processor = step.processor_class(input_path, output_path, **step.kwargs)
        error = processor.process()
        if error:
            print(f"Error in step {step.processor_class.__name__}: {error}")

    def _process_directory(self, step: ProcessingStep, input_dir: str, output_dir: str):
        if step.mode == 'subdirs':
            items = os.listdir(input_dir)
            for subdir in items:
                subdir_path = os.path.join(input_dir, subdir)
                if os.path.isdir(subdir_path):
                    subdir_output_path = os.path.join(output_dir, subdir)
                    os.makedirs(subdir_output_path, exist_ok=True)
                    self._process_file(step, subdir_path, subdir_output_path)
        elif step.mode == 'flat':
            self._process_flat_directory(step, input_dir, output_dir)
        elif step.mode == 'converge':
            # Special handling for converge mode
            pass  # Already handled in the main loop
        else:
            print(f"Error: Unsupported mode '{step.mode}' in step {step.processor_class.__name__}")

    def _process_flat_directory(self, step: ProcessingStep, input_dir: str, output_dir: str):
        for item in os.listdir(input_dir):
            item_path = os.path.join(input_dir, item)
            if os.path.isfile(item_path):
                item_output_dir = os.path.join(output_dir, os.path.splitext(item)[0])
                os.makedirs(item_output_dir, exist_ok=True)
                self._process_file(step, item_path, item_output_dir)
            elif os.path.isdir(item_path):
                print(f"Warning: Ignoring nested directory {item_path}")

"""
Example usage:

from extract_api import *

project = Project("docs","docs.zip","docs_out")
project.add_step(ZipExtractor)
project.add_step(DocxSplitter,regex_pattern="Answer:",keep_file_name=True)
project.add_step(FileCopier,'subdirs',file_extension='split_1.docx')
#project.add_step(PDFToImageConverter,'subdirs',combine_images=True)
project.add_step(Converger,'converge')
project.run()

"""
